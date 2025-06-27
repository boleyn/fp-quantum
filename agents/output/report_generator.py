"""
量子智能化功能点估算系统 - 报告生成器智能体

聚合估算结果，生成详细的分析报告和可视化图表
"""

import asyncio
from typing import Optional, List, Union
import logging
from datetime import datetime
import json
from pathlib import Path

from langchain_core.language_models import BaseLanguageModel
from langchain_openai import ChatOpenAI
from langchain_core.prompts import ChatPromptTemplate

from agents.base.base_agent import SpecializedAgent
from models.project_models import ProjectInfo
from models.nesma_models import NESMAEstimationResult, NESMAFunctionType
from models.cosmic_models import COSMICEstimationResult, COSMICDataMovementType
from models.common_models import (
    EstimationStrategy, VisualizationChart, ReportSection, 
    ExecutiveSummary, QualityAssessment, ReportData
)
from config.settings import get_settings

logger = logging.getLogger(__name__)


class ReportGeneratorAgent(SpecializedAgent):
    """报告生成器智能体 - 使用完整的Pydantic模型"""
    
    def __init__(self, llm: Optional[BaseLanguageModel] = None):
        super().__init__(
            agent_id="report_generator",
            specialty="report_generation",
            llm=llm
        )
        
        self.settings = get_settings()
        self.generated_reports: List[ReportData] = []
    
    def _get_capabilities(self) -> List[str]:
        """获取智能体能力列表"""
        return [
            "结果聚合分析",
            "多格式报告生成",
            "可视化图表创建",
            "标准对比报告",
            "质量评估报告"
        ]
    
    async def _execute_task(self, task_name: str, inputs: dict) -> dict:
        """执行报告生成任务"""
        if task_name == "generate_estimation_report":
            report_data = await self.generate_estimation_report(
                inputs["project_info"],
                inputs["estimation_results"],
                inputs.get("format", "markdown")
            )
            return report_data.dict()
        else:
            raise ValueError(f"未知任务: {task_name}")
    
    async def generate_estimation_report(
        self,
        project_info: ProjectInfo,
        estimation_results: dict,
        format: str = "markdown"
    ) -> ReportData:
        """生成完整的估算报告"""
        
        # 1. 创建执行摘要
        executive_summary = await self._create_executive_summary(
            estimation_results, project_info
        )
        
        # 2. 生成报告章节
        sections = await self._create_report_sections(
            estimation_results, project_info
        )
        
        # 3. 创建可视化数据
        visualizations = await self._create_visualizations(
            estimation_results
        )
        
        # 4. 生成质量评估
        quality_assessment = await self._create_quality_assessment(
            estimation_results
        )
        
        # 5. 生成报告内容
        content = await self._generate_report_content(
            executive_summary, sections, visualizations, quality_assessment, format
        )
        
        # 6. 创建完整报告数据
        report_data = ReportData(
            format=format,
            executive_summary=executive_summary,
            sections=sections,
            visualizations=visualizations,
            quality_assessment=quality_assessment,
            content=content,
            metadata={
            "project_name": project_info.name,
                "generation_time": datetime.now().isoformat(),
                "generator_version": "1.0"
            }
        )
        
        # 7. 记录生成历史
        self.generated_reports.append(report_data)
        
        return report_data
    
    async def _create_executive_summary(
        self,
        estimation_results: dict,
        project_info: ProjectInfo
    ) -> ExecutiveSummary:
        """创建执行摘要"""
        
        # 安全提取NESMA和COSMIC结果
        nesma_results = estimation_results.get("nesma_results")
        cosmic_results = estimation_results.get("cosmic_results")
        
        # 提取总功能点
        total_fp = self._extract_total_fp(nesma_results, cosmic_results)
        
        # 计算置信度
        confidence_level = self._calculate_overall_confidence(nesma_results, cosmic_results)
        
        # 识别使用的方法
        methods_used = self._identify_methods_used(nesma_results, cosmic_results)
        
        # 生成关键发现
        key_findings = await self._generate_key_findings(nesma_results, cosmic_results, project_info)
        
        # 确定估算策略
        strategy = self._determine_strategy(nesma_results, cosmic_results)
        
        # 生成格式化内容
        content = f"""
## 📊 执行摘要

### 项目概览
- **项目名称**: {project_info.name}
- **业务领域**: {project_info.business_domain}
- **估算日期**: {datetime.now().strftime("%Y-%m-%d")}
- **估算策略**: {strategy.value}

### 核心结果
- **总功能点**: {total_fp}
- **置信度**: {confidence_level}
- **估算方法**: {", ".join(methods_used)}

### 关键发现
{chr(10).join(f"- {finding}" for finding in key_findings)}
"""
        
        return ExecutiveSummary(
            project_name=project_info.name,
            business_domain=project_info.business_domain,
            estimation_date=datetime.now().strftime("%Y-%m-%d"),
            estimation_strategy=strategy,
            total_fp=total_fp,
            confidence_level=confidence_level,
            methods_used=methods_used,
            key_findings=key_findings,
            content=content
        )
    
    async def _create_report_sections(
        self,
        estimation_results: dict,
        project_info: ProjectInfo
    ) -> List[ReportSection]:
        """创建报告章节"""
        
        sections = []
        
        # NESMA详细分析章节
        nesma_results = estimation_results.get("nesma_results")
        if nesma_results:
            nesma_section = await self._create_nesma_section(nesma_results)
            sections.append(nesma_section)
        
        # COSMIC详细分析章节
        cosmic_results = estimation_results.get("cosmic_results")
        if cosmic_results:
            cosmic_section = await self._create_cosmic_section(cosmic_results)
            sections.append(cosmic_section)
        
        # 对比分析章节（如果两种标准都有结果）
        if nesma_results and cosmic_results:
            comparison_section = await self._create_comparison_section(nesma_results, cosmic_results)
            sections.append(comparison_section)
        
        return sections
    
    async def _create_nesma_section(self, nesma_results: dict) -> ReportSection:
        """创建NESMA章节"""
        
        # 安全解析NESMA结果
        total_ufp = nesma_results.get("total_ufp", 0)
        classifications = nesma_results.get("function_classifications", [])
        
        # 统计功能类型分布
        type_distribution = {}
        for classification in classifications:
            if classification is not None:
                if hasattr(classification, 'function_type'):
                    func_type = classification.function_type.value if hasattr(classification.function_type, 'value') else str(classification.function_type)
                else:
                    func_type = "Unknown"
                type_distribution[func_type] = type_distribution.get(func_type, 0) + 1
        
        content = f"""
## 🎯 NESMA 功能点分析

### 总体结果
- **未调整功能点(UFP)**: {total_ufp}
- **功能分类数量**: {len(classifications)}

### 功能类型分布
{chr(10).join(f"- **{func_type}**: {count}个" for func_type, count in type_distribution.items())}

### 详细分析
{self._format_nesma_details(classifications)}
"""
        
        return ReportSection(
            section_id="nesma_analysis",
            title="NESMA 功能点分析",
            content=content,
            order=1
        )
    
    async def _create_cosmic_section(self, cosmic_results: dict) -> ReportSection:
        """创建COSMIC章节"""
        
        # 安全解析COSMIC结果
        total_cfp = cosmic_results.get("total_cfp", 0)
        data_movements = cosmic_results.get("data_movements", [])
        
        # 统计数据移动类型分布
        movement_distribution = {}
        for movement in data_movements:
            if movement is not None:
                if hasattr(movement, 'movement_type'):
                    movement_type = movement.movement_type.value if hasattr(movement.movement_type, 'value') else str(movement.movement_type)
                else:
                    movement_type = "Unknown"
                movement_distribution[movement_type] = movement_distribution.get(movement_type, 0) + 1
        
        content = f"""
## 🌌 COSMIC 功能点分析

### 总体结果
- **COSMIC功能点(CFP)**: {total_cfp}
- **数据移动数量**: {len(data_movements)}

### 数据移动类型分布
{chr(10).join(f"- **{movement_type}**: {count}个" for movement_type, count in movement_distribution.items())}

### 详细分析
{self._format_cosmic_details(data_movements)}
"""
        
        return ReportSection(
            section_id="cosmic_analysis",
            title="COSMIC 功能点分析",
            content=content,
            order=2
        )
    
    async def _create_comparison_section(self, nesma_results: dict, cosmic_results: dict) -> ReportSection:
        """创建对比分析章节"""
        
        nesma_ufp = nesma_results.get("total_ufp", 0)
        cosmic_cfp = cosmic_results.get("total_cfp", 0)
        
        # 计算差异
        if nesma_ufp > 0 and cosmic_cfp > 0:
            difference = abs(nesma_ufp - cosmic_cfp)
            percentage = (difference / max(nesma_ufp, cosmic_cfp)) * 100
        else:
            difference = "N/A"
            percentage = "N/A"
        
        content = f"""
## ⚖️ 标准对比分析

### 结果对比
| 指标 | NESMA | COSMIC | 差异 |
|------|-------|---------|------|
| 总功能点 | {nesma_ufp} UFP | {cosmic_cfp} CFP | {difference} |
| 差异百分比 | - | - | {percentage}% |

### 差异分析
{self._analyze_differences(nesma_results, cosmic_results)}

### 推荐建议
{self._generate_recommendations(nesma_results, cosmic_results)}
"""
        
        return ReportSection(
            section_id="comparison_analysis",
            title="标准对比分析",
            content=content,
            order=3
        )
    
    async def _create_visualizations(self, estimation_results: dict) -> List[VisualizationChart]:
        """创建可视化图表"""
        
        visualizations = []
        
        # NESMA功能类型分布图
        nesma_results = estimation_results.get("nesma_results")
        if nesma_results:
            nesma_chart = self._create_nesma_distribution_chart(nesma_results)
            visualizations.append(nesma_chart)
        
        # COSMIC数据移动类型分布图
        cosmic_results = estimation_results.get("cosmic_results")
        if cosmic_results:
            cosmic_chart = self._create_cosmic_distribution_chart(cosmic_results)
            visualizations.append(cosmic_chart)
        
        return visualizations
    
    def _create_nesma_distribution_chart(self, nesma_results: dict) -> VisualizationChart:
        """创建NESMA分布图表"""
        
        classifications = nesma_results.get("function_classifications", [])
        
        # 统计各类型数量
        type_counts = {}
        for classification in classifications:
            if classification is not None:
                if hasattr(classification, 'function_type'):
                    func_type = classification.function_type.value if hasattr(classification.function_type, 'value') else str(classification.function_type)
                else:
                    func_type = "Unknown"
                type_counts[func_type] = type_counts.get(func_type, 0) + 1
        
        return VisualizationChart(
            chart_type="pie",
            title="NESMA功能类型分布",
            labels=list(type_counts.keys()),
            data=list(type_counts.values())
        )
    
    def _create_cosmic_distribution_chart(self, cosmic_results: dict) -> VisualizationChart:
        """创建COSMIC分布图表"""
        
        data_movements = cosmic_results.get("data_movements", [])
        
        # 统计各类型数量
        type_counts = {}
        for movement in data_movements:
            if movement is not None:
                if hasattr(movement, 'movement_type'):
                    movement_type = movement.movement_type.value if hasattr(movement.movement_type, 'value') else str(movement.movement_type)
                else:
                    movement_type = "Unknown"
                type_counts[movement_type] = type_counts.get(movement_type, 0) + 1
        
        return VisualizationChart(
            chart_type="pie",
            title="COSMIC数据移动类型分布",
            labels=list(type_counts.keys()),
            data=list(type_counts.values())
        )
    
    async def _create_quality_assessment(self, estimation_results: dict) -> QualityAssessment:
        """创建质量评估"""
        
        # 评估数据完整性
        completeness_score = self._assess_data_completeness(estimation_results)
        
        # 评估一致性
        consistency_score = self._assess_data_consistency(estimation_results)
        
        # 评估可信度
        reliability_score = self._assess_reliability(estimation_results)
        
        # 计算总体评分
        overall_score = (completeness_score + consistency_score + reliability_score) / 3
        
        # 生成建议
        recommendations = self._generate_quality_recommendations(
            completeness_score, consistency_score, reliability_score
        )
        
        return QualityAssessment(
            completeness_score=completeness_score,
            consistency_score=consistency_score,
            reliability_score=reliability_score,
            overall_score=overall_score,
            recommendations=recommendations
        )
    
    async def _generate_report_content(
        self,
        executive_summary: ExecutiveSummary,
        sections: List[ReportSection],
        visualizations: List[VisualizationChart],
        quality_assessment: QualityAssessment,
        format: str
    ) -> str:
        """生成报告内容"""
        
        if format.lower() == "markdown":
            return await self._create_markdown_content(
                executive_summary, sections, visualizations, quality_assessment
            )
        elif format.lower() == "excel":
            return await self._create_excel_content(
                executive_summary, sections, visualizations, quality_assessment
            )
        elif format.lower() == "word":
            return await self._create_word_content(
                executive_summary, sections, visualizations, quality_assessment
            )
        else:
            # 其他格式的实现
            return f"报告格式 {format} 暂未实现"
    
    async def _create_markdown_content(
        self,
        executive_summary: ExecutiveSummary,
        sections: List[ReportSection],
        visualizations: List[VisualizationChart],
        quality_assessment: QualityAssessment
    ) -> str:
        """创建Markdown格式内容"""
        
        content_parts = [
            "# 📊 功能点估算报告\n",
            executive_summary.content,
            "\n"
        ]
        
        # 添加各个章节
        for section in sorted(sections, key=lambda x: x.order):
            content_parts.append(section.content)
            content_parts.append("\n")
        
        # 添加质量评估
        content_parts.append(f"""
## 📈 质量评估

- **完整性评分**: {quality_assessment.completeness_score:.2f}
- **一致性评分**: {quality_assessment.consistency_score:.2f}
- **可信度评分**: {quality_assessment.reliability_score:.2f}
- **总体评分**: {quality_assessment.overall_score:.2f}

### 质量建议
{chr(10).join(f"- {rec}" for rec in quality_assessment.recommendations)}
""")
        
        # 添加可视化描述
        if visualizations:
            content_parts.append(f"""
## 📊 可视化图表

本报告包含 {len(visualizations)} 个图表：
{chr(10).join(f"- {chart.title}" for chart in visualizations)}
""")
        
        return "".join(content_parts)
    
    async def _create_excel_content(
        self,
        executive_summary: ExecutiveSummary,
        sections: List[ReportSection],
        visualizations: List[VisualizationChart],
        quality_assessment: QualityAssessment
    ) -> str:
        try:
            import pandas as pd
            from openpyxl import Workbook
            from openpyxl.styles import Font
            import tempfile
            import os
            
            wb = Workbook()
            wb.remove(wb.active)
            # 1. 执行摘要
            ws_summary = wb.create_sheet("执行摘要")
            ws_summary['A1'] = "功能点估算报告 - 执行摘要"
            ws_summary['A1'].font = Font(bold=True, size=16)
            summary_data = [
                ["项目名称", executive_summary.project_name],
                ["业务领域", executive_summary.business_domain],
                ["估算日期", executive_summary.estimation_date],
                ["估算策略", str(executive_summary.estimation_strategy)],
                ["总功能点", executive_summary.total_fp],
                ["置信度", executive_summary.confidence_level],
                ["估算方法", ", ".join(executive_summary.methods_used)]
            ]
            for i, (key, value) in enumerate(summary_data, 2):
                ws_summary[f'A{i}'] = key
                ws_summary[f'B{i}'] = str(value)
                ws_summary[f'A{i}'].font = Font(bold=True)
            ws_summary['A10'] = "关键发现"
            ws_summary['A10'].font = Font(bold=True, size=14)
            for i, finding in enumerate(executive_summary.key_findings, 11):
                ws_summary[f'A{i}'] = f"• {finding}"
            # 2. 各章节内容
            for section in sorted(sections, key=lambda x: x.order):
                ws = wb.create_sheet(section.title[:20])
                ws['A1'] = section.title
                ws['A1'].font = Font(bold=True, size=16)
                ws['A2'] = section.content
            # 3. 质量评估
            ws_quality = wb.create_sheet("质量评估")
            ws_quality['A1'] = "质量评估报告"
            ws_quality['A1'].font = Font(bold=True, size=16)
            quality_data = [
                ["完整性评分", f"{quality_assessment.completeness_score:.2f}"],
                ["一致性评分", f"{quality_assessment.consistency_score:.2f}"],
                ["可信度评分", f"{quality_assessment.reliability_score:.2f}"],
                ["总体评分", f"{quality_assessment.overall_score:.2f}"]
            ]
            for i, (key, value) in enumerate(quality_data, 3):
                ws_quality[f'A{i}'] = key
                ws_quality[f'B{i}'] = value
                ws_quality[f'A{i}'].font = Font(bold=True)
            ws_quality['A8'] = "质量建议"
            ws_quality['A8'].font = Font(bold=True, size=14)
            for i, rec in enumerate(quality_assessment.recommendations, 9):
                ws_quality[f'A{i}'] = f"• {rec}"
            # 4. 可视化
            if visualizations:
                ws_vis = wb.create_sheet("可视化图表")
                ws_vis['A1'] = "图表标题"
                ws_vis['B1'] = "数据"
                for idx, chart in enumerate(visualizations, 2):
                    ws_vis[f'A{idx}'] = chart.title
                    ws_vis[f'B{idx}'] = str(chart.data)
            # 保存文件
            with tempfile.NamedTemporaryFile(delete=False, suffix='.xlsx') as tmp:
                wb.save(tmp.name)
                return tmp.name
        except ImportError:
            return "生成Excel报告需要安装openpyxl库: uv pip install openpyxl"
        except Exception as e:
            return f"生成Excel报告失败: {str(e)}"
    
    async def _create_word_content(
        self,
        executive_summary: ExecutiveSummary,
        sections: List[ReportSection],
        visualizations: List[VisualizationChart],
        quality_assessment: QualityAssessment
    ) -> str:
        try:
            from docx import Document
            from docx.shared import Pt
            import tempfile
            import os
            doc = Document()
            doc.add_heading('功能点估算报告', 0)
            # 1. 执行摘要
            doc.add_heading('执行摘要', level=1)
            table = doc.add_table(rows=1, cols=2)
            table.style = 'Table Grid'
            hdr_cells = table.rows[0].cells
            hdr_cells[0].text = '项目信息'
            hdr_cells[1].text = '详情'
            project_info = [
                ['项目名称', executive_summary.project_name],
                ['业务领域', executive_summary.business_domain],
                ['估算日期', executive_summary.estimation_date],
                ['估算策略', str(executive_summary.estimation_strategy)],
                ['总功能点', str(executive_summary.total_fp)],
                ['置信度', executive_summary.confidence_level],
                ['估算方法', ', '.join(executive_summary.methods_used)]
            ]
            for info in project_info:
                row_cells = table.add_row().cells
                row_cells[0].text = info[0]
                row_cells[1].text = info[1]
            doc.add_heading('关键发现', level=2)
            for finding in executive_summary.key_findings:
                doc.add_paragraph(finding, style='List Bullet')
            # 2. 各章节内容
            for section in sorted(sections, key=lambda x: x.order):
                doc.add_heading(section.title, level=1)
                doc.add_paragraph(section.content)
            # 3. 质量评估
            doc.add_heading('质量评估', level=1)
            quality_table = doc.add_table(rows=1, cols=2)
            quality_table.style = 'Table Grid'
            quality_hdr = quality_table.rows[0].cells
            quality_hdr[0].text = '评估维度'
            quality_hdr[1].text = '评分'
            quality_data = [
                ['完整性评分', f"{quality_assessment.completeness_score:.2f}"],
                ['一致性评分', f"{quality_assessment.consistency_score:.2f}"],
                ['可信度评分', f"{quality_assessment.reliability_score:.2f}"],
                ['总体评分', f"{quality_assessment.overall_score:.2f}"]
            ]
            for data in quality_data:
                row_cells = quality_table.add_row().cells
                row_cells[0].text = data[0]
                row_cells[1].text = data[1]
            doc.add_heading('质量建议', level=2)
            for rec in quality_assessment.recommendations:
                doc.add_paragraph(rec, style='List Bullet')
            # 4. 可视化
            if visualizations:
                doc.add_heading('可视化图表', level=1)
                for chart in visualizations:
                    doc.add_paragraph(f"• {chart.title}: {str(chart.data)}")
            # 保存文件
            with tempfile.NamedTemporaryFile(delete=False, suffix='.docx') as tmp:
                doc.save(tmp.name)
                return tmp.name
        except ImportError:
            return "生成Word报告需要安装python-docx库: uv pip install python-docx"
        except Exception as e:
            return f"生成Word报告失败: {str(e)}"
    
    # 辅助方法
    def _extract_total_fp(self, nesma_results: dict, cosmic_results: dict) -> Union[int, str]:
        """提取总功能点数"""
        if nesma_results and cosmic_results:
            nesma_fp = nesma_results.get("total_ufp", 0)
            cosmic_fp = cosmic_results.get("total_cfp", 0)
            return f"NESMA: {nesma_fp}, COSMIC: {cosmic_fp}"
        elif nesma_results:
            return nesma_results.get("total_ufp", 0)
        elif cosmic_results:
            return cosmic_results.get("total_cfp", 0)
        else:
            return "未知"
    
    def _calculate_overall_confidence(self, nesma_results: dict, cosmic_results: dict) -> str:
        """计算总体置信度"""
        confidence_scores = []
        
        if nesma_results:
            nesma_confidence = nesma_results.get("average_confidence", 0)
            confidence_scores.append(nesma_confidence)
        
        if cosmic_results:
            cosmic_confidence = cosmic_results.get("average_confidence", 0)
            confidence_scores.append(cosmic_confidence)
        
        if confidence_scores:
            avg_confidence = sum(confidence_scores) / len(confidence_scores)
            if avg_confidence >= 0.8:
                return "高"
            elif avg_confidence >= 0.6:
                return "中"
            else:
                return "低"
        
        return "未知"
    
    def _identify_methods_used(self, nesma_results: dict, cosmic_results: dict) -> List[str]:
        """识别使用的方法"""
        methods = []
        
        if nesma_results:
            methods.append("NESMA v2.3+")
        
        if cosmic_results:
            methods.append("COSMIC v4.0+")
        
        return methods or ["未知方法"]
    
    def _determine_strategy(self, nesma_results: dict, cosmic_results: dict) -> EstimationStrategy:
        """确定估算策略"""
        if nesma_results and cosmic_results:
            return EstimationStrategy.DUAL_PARALLEL
        elif nesma_results:
            return EstimationStrategy.NESMA_ONLY
        elif cosmic_results:
            return EstimationStrategy.COSMIC_ONLY
        else:
            return EstimationStrategy.NESMA_ONLY  # 默认值
    
    async def _generate_key_findings(
        self, nesma_results: dict, cosmic_results: dict, project_info: ProjectInfo
    ) -> List[str]:
        """生成关键发现"""
        
        findings = []
        
        # 分析NESMA结果
        if nesma_results:
            classifications = nesma_results.get("function_classifications", [])
            if classifications:
                valid_classifications = [c for c in classifications if c is not None]
                if valid_classifications:
                    # 统计最常见的功能类型
                    type_counts = {}
                    for classification in valid_classifications:
                        if hasattr(classification, 'function_type'):
                            func_type = classification.function_type.value if hasattr(classification.function_type, 'value') else str(classification.function_type)
                            type_counts[func_type] = type_counts.get(func_type, 0) + 1
                    
                    if type_counts:
                        most_common_type = max(type_counts, key=type_counts.get)
                        findings.append(f"最常见的NESMA功能类型是{most_common_type}（{type_counts[most_common_type]}个）")
        
        # 分析COSMIC结果
        if cosmic_results:
            data_movements = cosmic_results.get("data_movements", [])
            if data_movements:
                valid_movements = [m for m in data_movements if m is not None]
                if valid_movements:
                    # 统计数据移动类型
                    movement_counts = {}
                    for movement in valid_movements:
                        if hasattr(movement, 'movement_type'):
                            movement_type = movement.movement_type.value if hasattr(movement.movement_type, 'value') else str(movement.movement_type)
                            movement_counts[movement_type] = movement_counts.get(movement_type, 0) + 1
                    
                    if movement_counts:
                        most_common_movement = max(movement_counts, key=movement_counts.get)
                        findings.append(f"最常见的COSMIC数据移动类型是{most_common_movement}（{movement_counts[most_common_movement]}个）")
        
        return findings if findings else ["标准的功能点估算结果"]
    
    def _format_nesma_details(self, classifications: List) -> str:
        """格式化NESMA详细信息"""
        if not classifications:
            return "无分类数据"
        
        details = []
        for i, classification in enumerate(classifications[:5], 1):  # 只显示前5个
            if classification is not None:
                if hasattr(classification, 'function_name'):
                    name = classification.function_name
                    func_type = classification.function_type.value if hasattr(classification.function_type, 'value') else str(classification.function_type)
                    details.append(f"{i}. **{name}** - {func_type}")
                else:
                    details.append(f"{i}. 分类数据格式异常")
        
        if len(classifications) > 5:
            details.append(f"... 共{len(classifications)}个功能")
        
        return "\n".join(details)
    
    def _format_cosmic_details(self, data_movements: List) -> str:
        """格式化COSMIC详细信息"""
        if not data_movements:
            return "无数据移动"
        
        details = []
        for i, movement in enumerate(data_movements[:5], 1):  # 只显示前5个
            if movement is not None:
                if hasattr(movement, 'data_group'):
                    data_group = movement.data_group
                    movement_type = movement.movement_type.value if hasattr(movement.movement_type, 'value') else str(movement.movement_type)
                    details.append(f"{i}. **{data_group}** - {movement_type}")
                else:
                    details.append(f"{i}. 数据移动格式异常")
        
        if len(data_movements) > 5:
            details.append(f"... 共{len(data_movements)}个数据移动")
        
        return "\n".join(details)
    
    def _analyze_differences(self, nesma_results: dict, cosmic_results: dict) -> str:
        """分析差异"""
        return "两种标准在估算方法和关注点上存在差异，NESMA侧重于功能复杂度，COSMIC侧重于数据移动。"
    
    def _generate_recommendations(self, nesma_results: dict, cosmic_results: dict) -> str:
        """生成推荐建议"""
        return "建议综合考虑两种标准的结果，结合项目实际情况选择最适合的估算标准。"
    
    def _assess_data_completeness(self, estimation_results: dict) -> float:
        """评估数据完整性"""
        score = 0.0
        
        if estimation_results.get("nesma_results"):
            score += 0.5
        if estimation_results.get("cosmic_results"):
            score += 0.5
        
        return min(score, 1.0)
    
    def _assess_data_consistency(self, estimation_results: dict) -> float:
        """评估数据一致性"""
        # 简单的一致性评估逻辑
        return 0.8  # 默认较高的一致性评分
    
    def _assess_reliability(self, estimation_results: dict) -> float:
        """评估可信度"""
        # 简单的可信度评估逻辑
        return 0.7  # 默认中等的可信度评分
    
    def _generate_quality_recommendations(
        self, completeness: float, consistency: float, reliability: float
    ) -> List[str]:
        """生成质量建议"""
        recommendations = []
        
        if completeness < 0.8:
            recommendations.append("建议收集更多完整的需求信息")
        if consistency < 0.8:
            recommendations.append("建议检查数据一致性")
        if reliability < 0.8:
            recommendations.append("建议增加验证步骤提高可信度")
        
        return recommendations or ["估算质量良好"]
    
    def get_generation_history(self) -> List[ReportData]:
        """获取生成历史"""
        return self.generated_reports
    
    def get_generation_statistics(self) -> dict:
        """获取生成统计"""
        return {
            "total_reports": len(self.generated_reports),
            "formats": list(set(report.format for report in self.generated_reports)),
            "latest_generation": self.generated_reports[-1].generated_at if self.generated_reports else None
        }


async def create_report_generator(llm: Optional[BaseLanguageModel] = None) -> ReportGeneratorAgent:
    """创建报告生成器智能体实例"""
    if llm is None:
        llm = ChatOpenAI(model="gpt-4o-mini", temperature=0.1)
    
    return ReportGeneratorAgent(llm=llm)


if __name__ == "__main__":
    async def main():
        # 测试报告生成器
        generator = await create_report_generator()
        logger.info(f"报告生成器创建成功，支持的能力: {generator._get_capabilities()}")
    
    asyncio.run(main()) 